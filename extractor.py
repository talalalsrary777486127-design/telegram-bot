import re
import io
import csv

URL_PATTERN = re.compile(
    r'(?:https?://|www\.)[^\s<>"\')\],;]+',
    re.IGNORECASE
)

TELEGRAM_PATTERN = re.compile(
    r'https?://(?:t\.me|telegram\.me|telegram\.org)(?:/[^\s<>"\')*\]},;]*)?|'
    r'@[a-zA-Z0-9_]{5,}',
    re.IGNORECASE
)

WHATSAPP_PATTERN = re.compile(
    r'https?://(?:wa\.me|api\.whatsapp\.com|chat\.whatsapp\.com|whatsapp\.com)(?:/[^\s<>"\')*\]},;]*)?',
    re.IGNORECASE
)

TG_GROUP_PATTERN = re.compile(
    r'https?://(?:t\.me|telegram\.me)/(?:joinchat/|\+)[^\s<>"\')*\]},;]*',
    re.IGNORECASE
)

WA_GROUP_PATTERN = re.compile(
    r'https?://chat\.whatsapp\.com/[^\s<>"\')*\]},;]*',
    re.IGNORECASE
)

WA_DIRECT_PATTERN = re.compile(
    r'https?://(?:wa\.me|api\.whatsapp\.com)/[^\s<>"\')*\]},;]*',
    re.IGNORECASE
)

# Pages with fewer than this many characters are considered image-based
_OCR_CHAR_THRESHOLD = 50
# DPI used when rasterising pages for OCR
_OCR_DPI = 200


def clean_url(url: str) -> str:
    return url.rstrip('.,;:!?)\]\'\"')


def classify_url(url: str) -> str:
    if WHATSAPP_PATTERN.match(url):
        return "whatsapp"
    if TELEGRAM_PATTERN.match(url):
        return "telegram"
    return "other"


def get_subcategory(url: str, category: str) -> str:
    """Classify into group/channel/direct/other subcategory."""
    if category == "telegram":
        if TG_GROUP_PATTERN.match(url):
            return "group"
        return "channel"
    elif category == "whatsapp":
        if WA_GROUP_PATTERN.match(url):
            return "group"
        if WA_DIRECT_PATTERN.match(url):
            return "direct"
        return "group"
    return "other"


def extract_from_text(text: str) -> dict[str, list[str]]:
    raw_urls = URL_PATTERN.findall(text)
    mentions = re.findall(r'@[a-zA-Z0-9_]{5,}', text)

    telegram, whatsapp, other = set(), set(), set()

    for url in raw_urls:
        url = clean_url(url)
        if not url:
            continue
        cat = classify_url(url)
        if cat == "telegram":
            telegram.add(url)
        elif cat == "whatsapp":
            whatsapp.add(url)
        else:
            other.add(url)

    for mention in mentions:
        telegram.add(mention)

    return {
        "telegram": sorted(telegram),
        "whatsapp": sorted(whatsapp),
        "other": sorted(other),
    }


def extract_from_txt(content: bytes) -> dict[str, list[str]]:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="replace")
    return extract_from_text(text)


def _read_pdf_text(reader) -> tuple[list[str], list[int]]:
    """
    Extract text from each page using pypdf.
    Returns (page_texts, thin_page_indices) where thin pages likely need OCR.
    """
    page_texts: list[str] = []
    thin_pages: list[int] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        # Harvest hyperlinks embedded in annotations
        try:
            if page.annotations:
                for annot in page.annotations:
                    obj = annot.get_object()
                    if obj and "/URI" in obj:
                        text += " " + str(obj["/URI"])
        except Exception:
            pass
        page_texts.append(text)
        if len(text.strip()) < _OCR_CHAR_THRESHOLD:
            thin_pages.append(i)

    return page_texts, thin_pages


def _ocr_pages(content: bytes, page_indices: list[int]) -> dict[int, str]:
    """
    Rasterise the requested pages and run Tesseract OCR on each.
    Returns {page_index: ocr_text}.
    """
    from pdf2image import convert_from_bytes
    import pytesseract

    images = convert_from_bytes(content, dpi=_OCR_DPI)
    results: dict[int, str] = {}
    for idx in page_indices:
        if idx < len(images):
            try:
                results[idx] = pytesseract.image_to_string(images[idx], lang="eng")
            except Exception:
                results[idx] = ""
    return results


def extract_from_pdf(content: bytes) -> tuple[dict[str, list[str]], dict]:
    """
    Extract links from a PDF.

    Returns:
        (extracted, stats)
        - extracted: {"telegram": [...], "whatsapp": [...], "other": [...]}
        - stats: {"pages": N, "ocr_pages": N, "method": "normal"|"ocr"|"mixed",
                  "ocr_error": str|None}
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    total_pages = len(reader.pages)

    page_texts, thin_pages = _read_pdf_text(reader)

    ocr_pages_count = 0
    ocr_error: str | None = None
    method = "normal"

    if thin_pages:
        try:
            ocr_results = _ocr_pages(content, thin_pages)
            for idx, text in ocr_results.items():
                page_texts[idx] = text
            ocr_pages_count = len(thin_pages)
            method = "ocr" if len(thin_pages) == total_pages else "mixed"
        except Exception as e:
            ocr_error = str(e)

    full_text = "\n".join(page_texts)
    extracted = extract_from_text(full_text)

    stats = {
        "pages": total_pages,
        "ocr_pages": ocr_pages_count,
        "method": method,
        "ocr_error": ocr_error,
    }

    return extracted, stats


def extract_from_docx(content: bytes) -> dict[str, list[str]]:
    """Extract links from a .docx file — body text, tables, headers/footers, and hyperlinks."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    # Paragraph text (body)
    for para in doc.paragraphs:
        parts.append(para.text)

    # Table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    # Headers and footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                for para in hf.paragraphs:
                    parts.append(para.text)

    # Embedded hyperlinks (href attributes in relationships)
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            parts.append(rel.target_ref)

    combined = "\n".join(parts)
    return extract_from_text(combined)


def extract_from_csv(content: bytes) -> dict[str, list[str]]:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="replace")

    all_text = []
    try:
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            all_text.extend(row)
    except Exception:
        all_text = [text]

    combined = "\n".join(all_text)
    return extract_from_text(combined)


def total_count(extracted: dict[str, list[str]]) -> int:
    return sum(len(v) for v in extracted.values())
